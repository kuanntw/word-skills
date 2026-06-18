#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公司 Word 文件樣板產生器（create_docx.py）
==========================================

使用 python-docx 套件建立標準化公司文件樣板，支援：
  - A4 頁面設定、頁首頁尾黑線與頁碼
  - 封面頁、修訂紀錄表、目錄／表目錄／圖目錄
  - 自訂 Heading 1-9、內文、程式碼區塊、表格／圖標號、清單等樣式
  - 從 JSON 讀取內容或產出示範文件（--demo）

相依套件：python-docx
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from copy import deepcopy
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import _Cell, _Row  # noqa: F401 (used for type hints)

# ---------------------------------------------------------------------------
# 常數（Constants）
# ---------------------------------------------------------------------------

# 主要品牌色
PRIMARY_COLOR = RGBColor(0x00, 0x33, 0x99)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY_999 = RGBColor(0x99, 0x99, 0x99)
CODE_BG = "F2F2F2"
CODE_BG_COLOR = RGBColor(0xF2, 0xF2, 0xF2)

# 字型名稱
FONT_CN = "微軟正黑體"   # 中文字型
FONT_EN = "Calibri"       # 英文字型
FONT_CODE = "Courier New" # 程式碼字型

# 頁面邊界（cm）
MARGIN_CM = 1.8

# 清單前綴對照表（5 層）
LIST_PREFIXES: Dict[str, List[str]] = {
    "letter":  ["A.", "(A)", "[A]", "a.", "(a)"],
    "chinese": ["一.", "(一)", "[一]", "甲.", "(甲)"],
    "symbol":  ["■", "◆", "●", "□", "◇"],
    "number":  ["1.", "(1)", "[1]", "I.", "(I)"],
}

# 清單層級縮排（左縮排、懸掛縮排，單位 cm）
LIST_INDENT: List[Tuple[float, float]] = [
    (1.0, 0.7),  # 第 1 層
    (1.7, 0.7),  # 第 2 層
    (2.4, 0.7),  # 第 3 層
    (3.1, 0.7),  # 第 4 層
    (3.8, 0.7),  # 第 5 層
]


# ---------------------------------------------------------------------------
# 輔助函式：OXML 與字型
# ---------------------------------------------------------------------------

def _set_run_font(
    run: Any,
    font_en: str = FONT_EN,
    font_cn: str = FONT_CN,
) -> None:
    """設定 run 的英文字型與中文字型。

    Args:
        run: python-docx Run 物件。
        font_en: 英文字型名稱。
        font_cn: 中文字型名稱（對應到 eastAsia）。
    """
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)


def _set_paragraph_spacing(
    paragraph: Any,
    before_pt: float = 0,
    after_pt: float = 0,
    line_spacing: Optional[float] = None,
    line_spacing_rule: int = WD_LINE_SPACING.SINGLE,
) -> None:
    """設定段落間距與行距。

    Args:
        paragraph: python-docx Paragraph 物件。
        before_pt: 段前間距（pt）。
        after_pt: 段後間距（pt）。
        line_spacing: 行距值（pt，用於固定行距時）。
        line_spacing_rule: 行距規則（WD_LINE_SPACING 列舉）。
    """
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)
    else:
        pf.line_spacing_rule = line_spacing_rule


def _add_field_code(paragraph: Any, field_code: str) -> None:
    """在段落中加入 OXML field code（例如 PAGE／TOC／NUMPAGES）。

    Args:
        paragraph: 目標 Paragraph。
        field_code: field code 字串，如 'PAGE' 或 'TOC \\o "1-3" \\h \\z \\u'。
    """
    run = paragraph.add_run()

    # w:fldChar begin
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar_begin)

    # w:instrText
    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_code
    run2._element.append(instrText)

    # w:fldChar separate
    run3 = paragraph.add_run()
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run3._element.append(fldChar_sep)

    # w:fldChar end
    run4 = paragraph.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run4._element.append(fldChar_end)


def _add_simple_page_number(paragraph: Any) -> None:
    """在段落中加入簡單的 PAGE field（頁碼）。"""
    _add_field_code(paragraph, " PAGE ")


def _add_toc_field(paragraph: Any, field_code: str) -> None:
    """建立完整的 TOC field，包含 begin／instrText／separate／end。

    Args:
        paragraph: 目標 Paragraph。
        field_code: TOC 指令字串。
    """
    _add_field_code(paragraph, field_code)


def _set_cell_shading(cell: Any, color_hex: str) -> None:
    """設定表格儲存格背景填充色。

    Args:
        cell: python-docx _Cell 物件。
        color_hex: 六位十六進位顏色碼（不含 #），例如 "F2F2F2"。
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    tcPr.append(shading)


def _set_cell_vertical_alignment(cell: Any, align: str = "center") -> None:
    """設定表格儲存格垂直對齊。

    Args:
        cell: python-docx _Cell 物件。
        align: "top" | "center" | "bottom"。
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _add_paragraph_border(
    paragraph: Any,
    position: str = "bottom",
    color: str = "000000",
    size: str = "12",
    space: str = "1",
) -> None:
    """為段落加入邊框。

    Args:
        paragraph: 目標 Paragraph。
        position: 邊框位置（"bottom" | "top" | "left" | "right"）。
        color: 線條顏色（hex 不帶 #）。
        size: 線條大小（eighth-points，2.25pt = 18）。
        space: 線條與文字的距離（pt）。
    """
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    pBdr.append(border)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# 樣式建立
# ---------------------------------------------------------------------------

def _create_all_styles(document: Document) -> None:
    """在 document 中建立所有自訂樣式。

    包含：
      - Heading 1-9（使用內建樣式並修改屬性）
      - 內文樣式（_內文-縮排, _內文-無縮排, 註解）
      - 程式碼樣式（程式碼-有編號, 程式碼-無編號）
      - 表／圖標號與表格樣式
      - 清單樣式（4 種清單各 5 層 + 清單內文 5 層 = 25 個）
    """
    # ---- Heading 1-9 ----
    heading_specs: List[Tuple[str, float, float, float, float]] = [
        ("Heading 1", 16, 12, 6, 0),
        ("Heading 2", 16, 10, 6, 0),
        ("Heading 3", 16, 10, 6, 0),
        ("Heading 4", 14, 8, 4, 0),
        ("Heading 5", 14, 8, 4, 0),
        ("Heading 6", 14, 8, 4, 0),
        ("Heading 7", 14, 8, 4, 0),
        ("Heading 8", 14, 8, 4, 0),
        ("Heading 9", 14, 8, 4, 0),
    ]

    for name, font_size, before, after, indent_left in heading_specs:
        # python-docx 內建的 heading 樣式需透過 style 物件修改
        try:
            style = document.styles[name]
        except KeyError:
            # 若不存在則新增（通常 Heading 1-9 內建存在）
            style = document.styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = PRIMARY_COLOR
        # 字型設定
        style.font.name = FONT_EN
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        # 段落格式
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if indent_left > 0:
            pf.left_indent = Cm(indent_left)
        # 設定 keep with next（避免標題孤行）
        pPr = style.element.get_or_add_pPr()
        keepNext = OxmlElement("w:keepNext")
        pPr.append(keepNext)

    # ---- 內文樣式 ----
    body_styles = [
        ("_內文-縮排", 12, BLACK, 1.0, 0.9, 2.5, 2.5, 20.0),
        ("_內文-無縮排", 12, BLACK, 1.0, 0.0, 2.5, 2.5, 20.0),
    ]
    for name, fs, color, left_indent, first_line, before, after, line_sp in body_styles:
        style = document.styles.add_style(name, 1)
        style.font.size = Pt(fs)
        style.font.color.rgb = color
        style.font.name = FONT_EN
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        pf = style.paragraph_format
        pf.left_indent = Cm(left_indent)
        pf.first_line_indent = Cm(first_line)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = Pt(line_sp)

    # ---- 註解樣式 ----
    note_style = document.styles.add_style("註解", 1)
    note_style.font.size = Pt(10.5)
    note_style.font.color.rgb = BLACK
    note_style.font.name = FONT_EN
    rPr = note_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    pf = note_style.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(0)

    # ---- 程式碼樣式 ----
    for code_name in ["程式碼-有編號", "程式碼-無編號"]:
        cs = document.styles.add_style(code_name, 1)
        cs.font.size = Pt(11)
        cs.font.name = FONT_CODE
        cs.font.color.rgb = BLACK
        rPr = cs.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:ascii"), FONT_CODE)
        rFonts.set(qn("w:hAnsi"), FONT_CODE)
        pf = cs.paragraph_format
        pf.left_indent = Cm(1.0)
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)

    # ---- 表／圖標號樣式 ----
    for cap_name in ["表_標號", "圖_標號"]:
        cap_style = document.styles.add_style(cap_name, 1)
        cap_style.font.size = Pt(12)
        cap_style.font.bold = True
        cap_style.font.name = FONT_CN
        rPr = cap_style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 表格標題樣式 ----
    tbl_title_style = document.styles.add_style("表格標題", 1)
    tbl_title_style.font.size = Pt(12)
    tbl_title_style.font.bold = True
    tbl_title_style.font.name = FONT_EN
    rPr = tbl_title_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    tbl_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 表格內容樣式 ----
    tbl_content_style = document.styles.add_style("表格內容", 1)
    tbl_content_style.font.size = Pt(12)
    tbl_content_style.font.name = FONT_EN
    rPr = tbl_content_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)

    # ---- 清單樣式（25 個） ----
    list_type_names = ["字母清單", "國字清單", "符號清單", "數字清單"]
    for type_name in list_type_names:
        for lv in range(1, 6):
            style_name = f"第{lv}層-{type_name}"
            ls = _create_list_style(document, style_name, lv)

    # ---- 清單內文樣式（5 個） ----
    for lv in range(1, 6):
        style_name = f"第{lv}層_____清單內文"
        ls_inner = document.styles.add_style(style_name, 1)
        ls_inner.font.size = Pt(12)
        ls_inner.font.name = FONT_EN
        rPr = ls_inner.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_CN)
        rFonts.set(qn("w:ascii"), FONT_EN)
        rFonts.set(qn("w:hAnsi"), FONT_EN)
        left_indent_cm, hanging_cm = LIST_INDENT[lv - 1]
        pf = ls_inner.paragraph_format
        pf.left_indent = Cm(left_indent_cm + hanging_cm)
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)


def _create_list_style(document: Document, style_name: str, level: int) -> Any:
    """建立單一清單層級樣式。

    Args:
        document: Document 物件。
        style_name: 樣式名稱。
        level: 層級（1-5）。

    Returns:
        建立的 Style 物件。
    """
    style = document.styles.add_style(style_name, 1)
    style.font.size = Pt(12)
    style.font.name = FONT_EN
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    left_indent_cm, hanging_cm = LIST_INDENT[level - 1]
    pf = style.paragraph_format
    pf.left_indent = Cm(left_indent_cm)
    # 懸掛縮排：首行負縮排
    pf.first_line_indent = Cm(-hanging_cm)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return style


# ---------------------------------------------------------------------------
# 頁面設定
# ---------------------------------------------------------------------------

def setup_page(document: Document) -> None:
    """設定 A4 直式頁面、邊界 1.8 cm。

    Args:
        document: Document 物件。
    """
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)


def add_header_footer(document: Document) -> None:
    """在每個 section 的頁首加入黑色粗橫線，頁尾加入黑色粗橫線與置中頁碼。

    Args:
        document: Document 物件。
    """
    for section in document.sections:
        # --- 頁首：黑色粗橫線（2.25 pt = 18 eighth-points） ---
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(4)
        _add_paragraph_border(hp, position="bottom", color="000000", size="18", space="4")

        # --- 頁尾：黑色粗橫線 + PAGE field ---
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after = Pt(0)
        _add_paragraph_border(fp, position="top", color="000000", size="18", space="4")
        _add_simple_page_number(fp)


# ---------------------------------------------------------------------------
# 核心函式
# ---------------------------------------------------------------------------

# --- 全域計數器（管理表／圖編號） ---
_table_counter: int = 0
_figure_counter: int = 0


def _reset_counters() -> None:
    """重設全域表／圖計數器。"""
    global _table_counter, _figure_counter
    _table_counter = 0
    _figure_counter = 0


def add_cover_page(
    document: Document,
    product_name: str,
    document_name: str,
    version: str,
    date: str,
) -> None:
    """建立封面頁。

    封面結構：
      - 中央偏上：產品名稱（26pt, 粗體, #003399）
                   文件名稱（20pt, 粗體, #003399）
                   版本（14pt, 黑色）
      - 底部中央：日期（12pt, 黑色）
      - 左下角：Template V1.1.0（8pt, 灰色 #999999）

    Args:
        document: Document 物件。
        product_name: 產品名稱。
        document_name: 文件名稱。
        version: 版本號。
        date: 日期字串。
    """
    # 清空第一頁現有內容
    if document.paragraphs:
        for p in document.paragraphs:
            p._element.getparent().remove(p._element)

    # 上半部空白區（將封面元素推至中央偏上）
    for _ in range(6):
        p = document.add_paragraph()
        p.text = ""
        _set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=12)

    # 產品名稱
    p_prod = document.add_paragraph()
    run_prod = p_prod.add_run(product_name)
    run_prod.font.size = Pt(26)
    run_prod.font.bold = True
    run_prod.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_prod)
    p_prod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_prod, before_pt=0, after_pt=8)

    # 文件名稱
    p_doc = document.add_paragraph()
    run_doc = p_doc.add_run(document_name)
    run_doc.font.size = Pt(20)
    run_doc.font.bold = True
    run_doc.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_doc)
    p_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_doc, before_pt=0, after_pt=8)

    # 版本
    p_ver = document.add_paragraph()
    run_ver = p_ver.add_run(version)
    run_ver.font.size = Pt(14)
    run_ver.font.color.rgb = BLACK
    _set_run_font(run_ver)
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_ver, before_pt=0, after_pt=4)

    # 中間空白區
    for _ in range(8):
        p = document.add_paragraph()
        p.text = ""
        _set_paragraph_spacing(p, before_pt=0, after_pt=0, line_spacing=12)

    # 日期（底部中央）
    p_date = document.add_paragraph()
    run_date = p_date.add_run(date)
    run_date.font.size = Pt(12)
    run_date.font.color.rgb = BLACK
    _set_run_font(run_date)
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p_date, before_pt=0, after_pt=12)

    # 左下角 Template 版本標示
    p_tmpl = document.add_paragraph()
    run_tmpl = p_tmpl.add_run("Template V1.1.0")
    run_tmpl.font.size = Pt(8)
    run_tmpl.font.color.rgb = GRAY_999
    _set_run_font(run_tmpl)
    p_tmpl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p_tmpl, before_pt=0, after_pt=0)

    # 封面後加入分節符號（下一頁）
    _add_section_break(document)


def add_revision_history(
    document: Document,
    revisions: List[Dict[str, str]],
) -> None:
    """建立修訂紀錄表。

    Args:
        document: Document 物件。
        revisions: 修訂紀錄 list，每筆為 dict，key: version / date / author / summary。
    """
    # 標題
    h = document.add_paragraph()
    run_h = h.add_run("修訂紀錄表")
    run_h.font.size = Pt(16)
    run_h.font.bold = True
    run_h.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_h)
    _set_paragraph_spacing(h, before_pt=12, after_pt=12)

    # 建立表格
    headers = ["版本", "日期", "作者", "修改摘要"]
    num_rows = len(revisions) + 1
    table = document.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表頭
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.size = Pt(12)
        run.font.bold = True
        _set_run_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "003399")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_vertical_alignment(cell, "center")

    # 資料列
    for row_idx, rev in enumerate(revisions):
        cells = table.rows[row_idx + 1].cells
        for col_idx, key in enumerate(["version", "date", "author", "summary"]):
            cell = cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(rev.get(key, ""))
            run.font.size = Pt(12)
            _set_run_font(run)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_vertical_alignment(cell, "center")

    # 表格後空行
    document.add_paragraph().text = ""

    # 分節符號
    _add_section_break(document)


def add_toc_page(document: Document) -> None:
    """建立目錄頁（含 TOC field code）。

    使用 OXML field: TOC \\o \"1-3\" \\h \\z \\u

    Args:
        document: Document 物件。
    """
    h = document.add_paragraph()
    run_h = h.add_run("目錄")
    run_h.font.size = Pt(20)
    run_h.font.bold = True
    run_h.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_h)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(h, before_pt=24, after_pt=12)

    # 空白行
    document.add_paragraph().text = ""

    # TOC field
    toc_p = document.add_paragraph()
    _add_toc_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u')

    # 分節符號
    _add_section_break(document)


def add_lot_page(document: Document) -> None:
    """建立表目錄頁。

    使用 TOC field: TOC \\h \\z \\c \"表\"

    Args:
        document: Document 物件。
    """
    h = document.add_paragraph()
    run_h = h.add_run("表目錄")
    run_h.font.size = Pt(20)
    run_h.font.bold = True
    run_h.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_h)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(h, before_pt=24, after_pt=12)

    document.add_paragraph().text = ""

    toc_p = document.add_paragraph()
    _add_toc_field(toc_p, 'TOC \\h \\z \\c "表"')

    _add_section_break(document)


def add_lof_page(document: Document) -> None:
    """建立圖目錄頁。

    使用 TOC field: TOC \\h \\z \\c \"圖\"

    Args:
        document: Document 物件。
    """
    h = document.add_paragraph()
    run_h = h.add_run("圖目錄")
    run_h.font.size = Pt(20)
    run_h.font.bold = True
    run_h.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_h)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(h, before_pt=24, after_pt=12)

    document.add_paragraph().text = ""

    toc_p = document.add_paragraph()
    _add_toc_field(toc_p, 'TOC \\h \\z \\c "圖"')

    _add_section_break(document)


def _add_section_break(document: Document) -> None:
    """加入分節符號（下一頁），並為新 section 設定相同頁面格式。

    Args:
        document: Document 物件。
    """
    new_section = document.add_section()
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.top_margin = Cm(MARGIN_CM)
    new_section.bottom_margin = Cm(MARGIN_CM)
    new_section.left_margin = Cm(MARGIN_CM)
    new_section.right_margin = Cm(MARGIN_CM)
    # 頁碼從 1 開始
    new_section.startING_page_number = 1


def add_code_block(
    document: Document,
    code: str,
    numbered: bool = False,
    language: str = "",
) -> None:
    """加入程式碼區塊。使用單欄表格實作灰底框線效果。

    Args:
        document: Document 物件。
        code: 程式碼字串（可含 \\n）。
        numbered: 是否加入行號。
        language: 程式語言標示（顯示於區塊上方）。
    """
    # 語言標籤
    if language:
        p_lang = document.add_paragraph()
        run_lang = p_lang.add_run(f"〔{language}〕")
        run_lang.font.size = Pt(10)
        run_lang.font.color.rgb = GRAY_999
        _set_run_font(run_lang)
        _set_paragraph_spacing(p_lang, before_pt=4, after_pt=0)

    # 建立單欄表格作為程式碼容器
    lines = code.split("\n")
    table = document.add_table(rows=len(lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 設定表格框線與寬度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    # 表格寬度 100%
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    # 移除已有 tblW
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)

    # 表格框線
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # 填入程式碼行
    for i, line in enumerate(lines):
        cell = table.rows[i].cells[0]
        cell.text = ""
        _set_cell_shading(cell, CODE_BG)
        _set_cell_vertical_alignment(cell, "center")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)

        # 行號
        if numbered:
            run_num = p.add_run(f"{i + 1:3d}  ")
            run_num.font.size = Pt(10)
            run_num.font.name = FONT_CODE
            run_num.font.color.rgb = GRAY_999

        run_code = p.add_run(line)
        run_code.font.size = Pt(11)
        run_code.font.name = FONT_CODE
        _set_run_font(run_code, font_en=FONT_CODE, font_cn=FONT_CN)

    # 區塊後空白
    document.add_paragraph().text = ""


def add_captioned_table(
    document: Document,
    title: str,
    headers: List[str],
    rows: List[List[str]],
    chapter_no: int = 1,
) -> None:
    """加入帶標號的表格。

    標號格式：表 1-1 表格名稱

    Args:
        document: Document 物件。
        title: 表格標題。
        headers: 表頭欄位名稱 list。
        rows: 資料列 list（每列為 list of str）。
        chapter_no: 章節編號（預設 1）。
    """
    global _table_counter
    _table_counter += 1

    # 表標號（使用 表_標號 樣式）
    caption_text = f"表 {chapter_no}-{_table_counter} {title}"
    p_cap = document.add_paragraph(style="表_標號")
    run_cap = p_cap.add_run(caption_text)
    _set_run_font(run_cap)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 建立表格
    num_rows = len(rows) + 1
    num_cols = len(headers)
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表頭
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.size = Pt(12)
        run.font.bold = True
        _set_run_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "003399")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_vertical_alignment(cell, "center")

    # 資料列
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(12)
            _set_run_font(run)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_vertical_alignment(cell, "center")

    # 表格後空行
    document.add_paragraph().text = ""


def add_captioned_image(
    document: Document,
    image_path: str,
    title: str,
    width: float = 5.5,
    chapter_no: int = 1,
) -> None:
    """加入帶標號的圖片。

    若無 image_path，則建立 placeholder 方塊（灰底虛線框）。

    Args:
        document: Document 物件。
        image_path: 圖片檔案路徑（空字串則用 placeholder）。
        title: 圖標題。
        width: 圖片寬度（英吋）。
        chapter_no: 章節編號。
    """
    global _figure_counter
    _figure_counter += 1

    # 圖標號（使用 圖_標號 樣式）
    caption_text = f"圖 {chapter_no}-{_figure_counter} {title}"
    p_cap = document.add_paragraph(style="圖_標號")
    run_cap = p_cap.add_run(caption_text)
    _set_run_font(run_cap)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 圖片或 placeholder
    if image_path and os.path.isfile(image_path):
        p_img = document.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        try:
            run_img.add_picture(image_path, width=Inches(width))
        except Exception:
            _add_image_placeholder(document, title, width)
    else:
        _add_image_placeholder(document, title, width)

    document.add_paragraph().text = ""


def _add_image_placeholder(document: Document, title: str, width: float) -> None:
    """建立圖片 placeholder（灰底虛線框 + 提示文字）。

    Args:
        document: Document 物件。
        title: 圖標題（用於提示文字）。
        width: placeholder 寬度（英吋，僅用於參考）。
    """
    # 使用單欄表格模擬圖片 placeholder
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    _set_cell_shading(cell, "F5F5F5")
    _set_cell_vertical_alignment(cell, "center")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"〔圖片 Placeholder: {title}〕")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY_999
    run.font.italic = True
    _set_run_font(run)
    _set_paragraph_spacing(p, before_pt=36, after_pt=36)
    # 設虛線框線
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "dashed")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_list_items(
    document: Document,
    list_type: str,
    items: List[Dict[str, Any]],
) -> None:
    """加入清單項目。

    list_type 支援: "letter" | "chinese" | "symbol" | "number"
    items 格式: [{"level": 1, "text": "...", "note": "..."}, ...]

    Args:
        document: Document 物件。
        list_type: 清單類型。
        items: 清單項目 list。
    """
    type_name_map = {
        "letter": "字母清單",
        "chinese": "國字清單",
        "symbol": "符號清單",
        "number": "數字清單",
    }
    type_cn = type_name_map.get(list_type, "數字清單")
    prefixes = LIST_PREFIXES.get(list_type, LIST_PREFIXES["number"])

    for item in items:
        level = max(1, min(5, item.get("level", 1)))
        text = item.get("text", "")
        note = item.get("note", "")

        prefix = prefixes[level - 1] if level <= len(prefixes) else ""
        full_text = f"{prefix} {text}"
        style_name = f"第{level}層-{type_cn}"

        # 確定樣式是否存在
        try:
            doc_style = document.styles[style_name]
        except KeyError:
            doc_style = None

        p = document.add_paragraph()
        if doc_style:
            p.style = doc_style
        run = p.add_run(full_text)
        run.font.size = Pt(12)
        _set_run_font(run)
        left_indent_cm, hanging_cm = LIST_INDENT[level - 1]
        p.paragraph_format.left_indent = Cm(left_indent_cm)
        p.paragraph_format.first_line_indent = Cm(-hanging_cm)
        _set_paragraph_spacing(p, before_pt=2, after_pt=2)

        # 若有附註（note），以清單內文樣式加入
        if note:
            inner_style_name = f"第{level}層_____清單內文"
            try:
                inner_style = document.styles[inner_style_name]
            except KeyError:
                inner_style = None
            p_note = document.add_paragraph()
            if inner_style:
                p_note.style = inner_style
            run_note = p_note.add_run(note)
            run_note.font.size = Pt(12)
            _set_run_font(run_note)
            inner_left = left_indent_cm + hanging_cm
            p_note.paragraph_format.left_indent = Cm(inner_left)
            p_note.paragraph_format.first_line_indent = Cm(0)
            _set_paragraph_spacing(p_note, before_pt=2, after_pt=2)


def add_appendix(
    document: Document,
    label: str,
    title: str,
    blocks: List[Dict[str, Any]],
) -> None:
    """加入附錄。

    Args:
        document: Document 物件。
        label: 附錄標籤（如 "A"、"B"）。
        title: 附錄標題。
        blocks: 附錄內容區塊 list（與 section blocks 格式相同）。
    """
    # 附錄標題（Heading 1 格式）
    h = document.add_paragraph()
    run_h = h.add_run(f"附錄 {label}：{title}")
    run_h.font.size = Pt(16)
    run_h.font.bold = True
    run_h.font.color.rgb = PRIMARY_COLOR
    _set_run_font(run_h)
    _set_paragraph_spacing(h, before_pt=12, after_pt=6)
    pPr = h._element.get_or_add_pPr()
    keepNext = OxmlElement("w:keepNext")
    pPr.append(keepNext)

    # 附錄內容
    _process_blocks(document, blocks, chapter_no=1)


def _process_blocks(
    document: Document,
    blocks: List[Dict[str, Any]],
    chapter_no: int = 1,
) -> None:
    """處理內容區塊 list（paragraph / code / table / image / list）。

    Args:
        document: Document 物件。
        blocks: 內容區塊 list。
        chapter_no: 當前章節編號（用於表／圖編號）。
    """
    for block in blocks:
        btype = block.get("type", "paragraph")

        if btype == "paragraph":
            text = block.get("text", "")
            indent = block.get("indent", True)
            p = document.add_paragraph()
            if indent:
                try:
                    p.style = document.styles["_內文-縮排"]
                except KeyError:
                    pass
            else:
                try:
                    p.style = document.styles["_內文-無縮排"]
                except KeyError:
                    pass
            run = p.add_run(text)
            run.font.size = Pt(12)
            _set_run_font(run)

        elif btype == "code":
            code_text = block.get("code", "")
            numbered = block.get("numbered", False)
            language = block.get("language", "")
            add_code_block(document, code_text, numbered=numbered, language=language)

        elif btype == "table":
            title = block.get("title", "")
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            add_captioned_table(document, title, headers, rows, chapter_no=chapter_no)

        elif btype == "image":
            title = block.get("title", "")
            path = block.get("path", "")
            width = block.get("width", 5.5)
            add_captioned_image(document, path, title, width=width, chapter_no=chapter_no)

        elif btype == "list":
            list_type = block.get("list_type", "number")
            items = block.get("items", [])
            add_list_items(document, list_type, items)


def add_section_content(
    document: Document,
    sections: List[Dict[str, Any]],
    heading_numbers: Optional[List[int]] = None,
) -> None:
    """遞迴加入章節內容。

    Args:
        document: Document 物件。
        sections: 章節定義 list。
        heading_numbers: 當前標題編號（用於遞迴追蹤），例如 [1, 2] 表示 1.2。
    """
    if heading_numbers is None:
        heading_numbers = []

    for section in sections:
        level = section.get("level", 1)
        title = section.get("title", "")
        blocks = section.get("blocks", [])
        children = section.get("children", [])

        # 計算標題編號
        if len(heading_numbers) >= level:
            heading_numbers = heading_numbers[: level - 1]
            heading_numbers.append(heading_numbers[-1] + 1 if heading_numbers else 1)
        else:
            while len(heading_numbers) < level:
                heading_numbers.append(1 if len(heading_numbers) == 0 else 1)
            heading_numbers[level - 1] = heading_numbers[level - 1]

        num_str = ".".join(str(n) for n in heading_numbers[:level])
        heading_text = f"{num_str} {title}"

        # 加入標題段落（使用 Heading 樣式）
        heading_style_name = f"Heading {level}"
        h = document.add_paragraph(style=heading_style_name)
        run_h = h.add_run(heading_text)
        _set_run_font(run_h)
        # keep with next
        pPr = h._element.get_or_add_pPr()
        keepNext = OxmlElement("w:keepNext")
        pPr.append(keepNext)

        # 處理內容區塊
        _process_blocks(document, blocks, chapter_no=heading_numbers[0] if heading_numbers else 1)

        # 遞迴處理子章節
        if children:
            add_section_content(document, children, deepcopy(heading_numbers))


# ---------------------------------------------------------------------------
# 文件建立入口
# ---------------------------------------------------------------------------

def create_demo_document(output_path: str) -> None:
    """建立完整示範文件，包含所有樣式與結構。

    Args:
        output_path: 輸出 .docx 檔案路徑。
    """
    document = Document()
    _reset_counters()

    # 頁面設定
    setup_page(document)

    # 建立所有自訂樣式
    _create_all_styles(document)

    # ---- 封面 ----
    add_cover_page(
        document,
        product_name="示範產品 DemoProduct",
        document_name="系統設計規格書",
        version="V1.0.0",
        date="2024/06/18",
    )

    # ---- 修訂紀錄表 ----
    revisions = [
        {"version": "V1.0.0", "date": "2024/06/18", "author": "文件作者", "summary": "初版建立"},
        {"version": "V0.5.0", "date": "2024/06/01", "author": "文件作者", "summary": "內部審查版本"},
    ]
    add_revision_history(document, revisions)

    # ---- 目錄頁 ----
    add_toc_page(document)

    # ---- 表目錄 ----
    add_lot_page(document)

    # ---- 圖目錄 ----
    add_lof_page(document)

    # ---- 正文（所有層級標題 + 內容） ----
    demo_sections: List[Dict[str, Any]] = [
        {
            "level": 1,
            "title": "文件目的",
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "本文件旨在展示公司標準文件樣板之完整功能，包含所有 Heading 層級、內文樣式、清單樣式、程式碼區塊、表格與圖片。",
                    "indent": True,
                },
                {
                    "type": "paragraph",
                    "text": "此為無縮排內文範例，用於不需要首行縮排的段落。",
                    "indent": False,
                },
            ],
            "children": [
                {
                    "level": 2,
                    "title": "適用範圍",
                    "blocks": [
                        {
                            "type": "paragraph",
                            "text": "本文件適用於所有開發團隊、測試團隊及專案管理相關人員。所有文件撰寫應遵循本樣板規範。",
                            "indent": True,
                        },
                    ],
                    "children": [
                        {
                            "level": 3,
                            "title": "參考文件",
                            "blocks": [
                                {
                                    "type": "paragraph",
                                    "text": "相關參考文件包含：系統需求規格書（SRS）、API 設計文件、資料庫設計文件。",
                                    "indent": True,
                                },
                            ],
                            "children": [
                                {
                                    "level": 4,
                                    "title": "外部標準",
                                    "blocks": [
                                        {
                                            "type": "paragraph",
                                            "text": "本文件參考 ISO 9001 品質管理標準與 CMMI Level 3 流程規範。",
                                            "indent": True,
                                        },
                                    ],
                                    "children": [
                                        {
                                            "level": 5,
                                            "title": "ISO 9001 章節對照",
                                            "blocks": [
                                                {
                                                    "type": "paragraph",
                                                    "text": "本文件之品質管理章節對應 ISO 9001:2015 第 8 章營運規劃與控制。",
                                                    "indent": True,
                                                },
                                            ],
                                            "children": [
                                                {
                                                    "level": 6,
                                                    "title": "文件管制要求",
                                                    "blocks": [
                                                        {
                                                            "type": "paragraph",
                                                            "text": "所有文件變更均需經過審查與核准程序。",
                                                            "indent": True,
                                                        },
                                                    ],
                                                    "children": [
                                                        {
                                                            "level": 7,
                                                            "title": "版本控制細則",
                                                            "blocks": [
                                                                {
                                                                    "type": "paragraph",
                                                                    "text": "版本號格式為 V<主版號>.<次版號>.<修訂號>，例如 V1.2.3。",
                                                                    "indent": True,
                                                                },
                                                            ],
                                                            "children": [
                                                                {
                                                                    "level": 8,
                                                                    "title": "版號遞增規則",
                                                                    "blocks": [
                                                                        {
                                                                            "type": "paragraph",
                                                                            "text": "主版號變更代表重大架構調整，次版號變更代表功能新增，修訂號變更代表錯誤修正。",
                                                                            "indent": True,
                                                                        },
                                                                    ],
                                                                    "children": [
                                                                        {
                                                                            "level": 9,
                                                                            "title": "自動化版號管理",
                                                                            "blocks": [
                                                                                {
                                                                                    "type": "paragraph",
                                                                                    "text": "建議使用 Git tag 搭配 CI/CD pipeline 自動管理版本號。",
                                                                                    "indent": True,
                                                                                },
                                                                            ],
                                                                            "children": [],
                                                                        },
                                                                    ],
                                                                },
                                                            ],
                                                        },
                                                    ],
                                                },
                                            ],
                                        },
                                    ],
                                },
                            ],
                        },
                    ],
                },
            ],
        },
        {
            "level": 1,
            "title": "功能需求",
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "以下列出系統主要功能需求，以符號清單方式呈現。",
                    "indent": True,
                },
                {
                    "type": "list",
                    "list_type": "symbol",
                    "items": [
                        {"level": 1, "text": "使用者註冊與登入功能"},
                        {"level": 1, "text": "訂單建立與查詢功能"},
                        {"level": 2, "text": "訂單狀態追蹤（處理中、已出貨、已完成）"},
                        {"level": 2, "text": "訂單歷史查詢（支援日期範圍篩選）"},
                        {"level": 1, "text": "報表匯出功能（支援 PDF 與 Excel 格式）"},
                    ],
                },
                {
                    "type": "paragraph",
                    "text": "以下為操作流程之數字清單範例（含清單內文 note）。",
                    "indent": True,
                },
                {
                    "type": "list",
                    "list_type": "number",
                    "items": [
                        {
                            "level": 1,
                            "text": "使用者開啟系統並登入",
                            "note": "系統將驗證帳號密碼，若啟用雙因子驗證則需輸入 OTP。",
                        },
                        {"level": 1, "text": "進入主控台後選擇所需功能模組"},
                        {
                            "level": 2,
                            "text": "選擇訂單管理模組",
                            "note": "訂單管理模組提供建立、查詢、修改、刪除等功能。",
                        },
                        {
                            "level": 3,
                            "text": "點選「建立新訂單」按鈕",
                        },
                        {
                            "level": 4,
                            "text": "填寫訂單詳細資料（品項、數量、運送地址）",
                        },
                        {
                            "level": 5,
                            "text": "確認無誤後點選「送出訂單」",
                            "note": "系統將自動發送確認通知至使用者 Email。",
                        },
                    ],
                },
                {
                    "type": "table",
                    "title": "功能需求一覽表",
                    "headers": ["項次", "功能名稱", "優先級", "負責模組", "狀態"],
                    "rows": [
                        ["1", "使用者註冊", "高", "User Module", "已完成"],
                        ["2", "使用者登入", "高", "Auth Module", "已完成"],
                        ["3", "訂單建立", "高", "Order Module", "開發中"],
                        ["4", "報表匯出", "中", "Report Module", "規劃中"],
                    ],
                },
                {
                    "type": "image",
                    "title": "系統功能模組架構圖",
                    "path": "",
                    "width": 5.5,
                },
            ],
            "children": [
                {
                    "level": 2,
                    "title": "分類清單範例（字母清單）",
                    "blocks": [
                        {
                            "type": "list",
                            "list_type": "letter",
                            "items": [
                                {"level": 1, "text": "前端技術分類"},
                                {"level": 2, "text": "React 生態系"},
                                {"level": 3, "text": "React Hooks 與 Context API"},
                                {"level": 1, "text": "後端技術分類"},
                                {"level": 2, "text": "Spring Boot 框架"},
                                {"level": 3, "text": "Spring Security 權限控制"},
                            ],
                        },
                    ],
                },
                {
                    "level": 2,
                    "title": "規範清單範例（國字清單）",
                    "blocks": [
                        {
                            "type": "list",
                            "list_type": "chinese",
                            "items": [
                                {"level": 1, "text": "總則"},
                                {"level": 2, "text": "適用範圍與定義"},
                                {"level": 3, "text": "本規範適用於全體開發人員與專案管理人員"},
                                {"level": 1, "text": "程式碼規範"},
                                {"level": 2, "text": "命名規則"},
                                {"level": 3, "text": "使用駝峰式命名法（camelCase）"},
                            ],
                        },
                    ],
                },
            ],
        },
        {
            "level": 1,
            "title": "程式碼範例",
            "blocks": [
                {
                    "type": "paragraph",
                    "text": "以下為有編號的 Python 程式碼範例：",
                    "indent": True,
                },
                {
                    "type": "code",
                    "numbered": True,
                    "language": "python",
                    "code": 'def hello_world(name: str) -> str:\n    """Greet the user."""\n    greeting = f"Hello, {name}!"\n    return greeting\n\n\nif __name__ == "__main__":\n    result = hello_world("World")\n    print(result)',
                },
                {
                    "type": "paragraph",
                    "text": "以下為無編號的 SQL 程式碼範例：",
                    "indent": True,
                },
                {
                    "type": "code",
                    "numbered": False,
                    "language": "sql",
                    "code": "SELECT u.username, o.order_date, o.total\nFROM users u\nJOIN orders o ON u.id = o.user_id\nWHERE o.status = 'completed'\nORDER BY o.order_date DESC;",
                },
            ],
            "children": [],
        },
    ]

    add_section_content(document, demo_sections)

    # ---- 附錄 ----
    appendix_a_blocks: List[Dict[str, Any]] = [
        {
            "type": "paragraph",
            "text": "本附錄補充說明系統所使用的資料庫結構。",
            "indent": True,
        },
        {
            "type": "code",
            "numbered": False,
            "language": "sql",
            "code": "CREATE TABLE products (\n    id SERIAL PRIMARY KEY,\n    name VARCHAR(200) NOT NULL,\n    price DECIMAL(10, 2) NOT NULL,\n    stock INT DEFAULT 0,\n    created_at TIMESTAMP DEFAULT NOW()\n);",
        },
        {
            "type": "table",
            "title": "products 資料表欄位說明",
            "headers": ["欄位名稱", "型別", "說明", "備註"],
            "rows": [
                ["id", "SERIAL", "主鍵，自動遞增", "PRIMARY KEY"],
                ["name", "VARCHAR(200)", "產品名稱", "NOT NULL"],
                ["price", "DECIMAL(10,2)", "產品價格", "NOT NULL"],
                ["stock", "INT", "庫存數量", "預設 0"],
                ["created_at", "TIMESTAMP", "建立時間", "預設 NOW()"],
            ],
        },
    ]

    appendix_b_blocks: List[Dict[str, Any]] = [
        {
            "type": "paragraph",
            "text": "本附錄列出系統所有錯誤碼及其對應說明。",
            "indent": True,
        },
        {
            "type": "table",
            "title": "系統錯誤碼對照表",
            "headers": ["錯誤碼", "HTTP Status", "說明", "處理建議"],
            "rows": [
                ["E001", "400", "請求參數錯誤", "請檢查輸入格式"],
                ["E002", "401", "未認證", "請重新登入"],
                ["E003", "403", "權限不足", "請聯絡管理員"],
                ["E004", "404", "資源不存在", "請確認資源 ID"],
                ["E005", "500", "伺服器內部錯誤", "請聯絡技術支援"],
            ],
        },
    ]

    add_appendix(document, "A", "資料庫結構補充", appendix_a_blocks)
    add_appendix(document, "B", "錯誤碼對照表", appendix_b_blocks)

    # ---- 加入頁首頁尾 ----
    add_header_footer(document)

    # ---- 存檔 ----
    output_dir = os.path.dirname(os.path.abspath(output_path))
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    document.save(output_path)
    print(f"[OK] 示範文件已儲存至：{os.path.abspath(output_path)}")


def create_document_from_json(json_path: str, output_path: str) -> None:
    """從 JSON 讀取內容並建立文件。

    JSON 結構範例請參考 example_input.json。

    Args:
        json_path: JSON 輸入檔案路徑。
        output_path: 輸出 .docx 檔案路徑。
    """
    # 讀取 JSON
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data: Dict[str, Any] = json.load(f)
    except FileNotFoundError:
        print(f"[ERROR] 找不到 JSON 檔案：{json_path}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"[ERROR] JSON 解析失敗：{e}", file=sys.stderr)
        sys.exit(1)

    product_name = data.get("product_name", "產品名稱")
    document_name = data.get("document_name", "文件名稱")
    version = data.get("version", "V1.0.0")
    date = data.get("date", "2024/01/01")
    revision_history = data.get("revision_history", [
        {"version": "V1.0.0", "date": date, "author": "文件作者", "summary": "初版建立"},
    ])
    sections = data.get("sections", [])
    appendices = data.get("appendices", [])

    document = Document()
    _reset_counters()

    # 頁面設定
    setup_page(document)

    # 建立樣式
    _create_all_styles(document)

    # 封面
    add_cover_page(document, product_name, document_name, version, date)

    # 修訂紀錄
    add_revision_history(document, revision_history)

    # 目錄
    add_toc_page(document)

    # 表目錄
    add_lot_page(document)

    # 圖目錄
    add_lof_page(document)

    # 正文
    add_section_content(document, sections)

    # 附錄
    for appendix in appendices:
        label = appendix.get("label", "")
        title = appendix.get("title", "")
        blocks = appendix.get("blocks", [])
        add_appendix(document, label, title, blocks)

    # 頁首頁尾
    add_header_footer(document)

    # 存檔
    output_dir = os.path.dirname(os.path.abspath(output_path))
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    document.save(output_path)
    print(f"[OK] 文件已儲存至：{os.path.abspath(output_path)}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    """CLI 進入點：使用 argparse 解析參數並執行對應功能。"""
    parser = argparse.ArgumentParser(
        description="公司 Word 文件樣板產生器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用範例：
  python create_docx.py --demo --output demo.docx
  python create_docx.py --input example_input.json --output output.docx
        """,
    )
    parser.add_argument(
        "--input",
        type=str,
        default=None,
        help="JSON 輸入檔案路徑",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="output.docx",
        help="輸出 .docx 路徑（預設：output.docx）",
    )
    parser.add_argument(
        "--demo",
        action="store_true",
        default=False,
        help="產出示範文件",
    )

    args = parser.parse_args()

    # 驗證參數
    if not args.demo and not args.input:
        parser.print_help()
        print("\n[ERROR] 請指定 --input 或 --demo", file=sys.stderr)
        sys.exit(1)

    try:
        if args.demo:
            create_demo_document(args.output)
        elif args.input:
            create_document_from_json(args.input, args.output)
    except Exception as e:
        print(f"[ERROR] 執行失敗：{e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()