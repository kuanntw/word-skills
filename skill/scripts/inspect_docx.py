#!/usr/bin/env python3
"""檢查產出的 Word 文件是否符合公司樣式規範（支援 .docx 與 .doc）。

用法：
    python inspect_docx.py output.docx
    python inspect_docx.py CM03A_V1.doc

檢查項目：
    .docx：樣式、錯誤參照、Normal 段落比例、封面、修訂紀錄、標題、表格標號、頁首頁尾
    .doc：錯誤參照、修訂紀錄、內容結構、基本格式（樣式檢查需先轉為 .docx）
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Tuple


# 必要樣式清單
REQUIRED_STYLES = [
    "Heading 1", "Heading 2", "Heading 3",
    "Heading 4", "Heading 5", "Heading 6",
    "Heading 7", "Heading 8", "Heading 9",
    "_內文-縮排",
    "_內文-無縮排",
    "註解",
    "程式碼-有編號",
    "程式碼-無編號",
    "表_標號",
    "圖_標號",
    "表格標題",
    "表格內容",
]

COVER_REQUIRED = [
    "Template",
]

LIST_STYLE_PREFIXES = [
    "第1層-字母清單", "第2層-字母清單", "第3層-字母清單",
    "第4層-字母清單", "第5層-字母清單",
    "第1層-國字清單", "第2層-國字清單", "第3層-國字清單",
    "第4層-國字清單", "第5層-國字清單",
    "第1層-符號清單", "第2層-符號清單", "第3層-符號清單",
    "第4層-符號清單", "第5層-符號清單",
    "第1層-數字清單", "第2層-數字清單", "第3層-數字清單",
    "第4層-數字清單", "第5層-數字清單",
    "第1層_____清單內文", "第2層_____清單內文", "第3層_____清單內文",
    "第4層_____清單內文", "第5層_____清單內文",
]


# 格式判斷
def _get_file_type(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext == '.docx':
        return 'docx'
    elif ext == '.doc':
        return 'doc'
    else:
        return 'unknown'


class InspectionResult:
    """檢查結果收集器。"""
    def __init__(self) -> None:
        self.passes: List[str] = []
        self.warnings: List[str] = []
        self.failures: List[str] = []

    def add_pass(self, msg: str) -> None:
        self.passes.append(msg)

    def add_warn(self, msg: str) -> None:
        self.warnings.append(msg)

    def add_fail(self, msg: str) -> None:
        self.failures.append(msg)

    def has_failures(self) -> bool:
        return len(self.failures) > 0

    def print_summary(self) -> None:
        total = len(self.passes) + len(self.warnings) + len(self.failures)
        print(f"\n{'='*60}")
        print(f"檢查完成：共 {total} 項檢查")
        print(f"  PASS: {len(self.passes)}")
        print(f"  WARN: {len(self.warnings)}")
        print(f"  FAIL: {len(self.failures)}")
        print(f"{'='*60}")

        if self.failures:
            print("\n--- FAIL 項目 ---")
            for f in self.failures:
                print(f"  [FAIL] {f}")

        if self.warnings:
            print("\n--- WARN 項目 ---")
            for w in self.warnings:
                print(f"  [WARN] {w}")

        if self.passes:
            print("\n--- PASS 項目 ---")
            for p in self.passes:
                print(f"  [PASS] {p}")

        if self.has_failures():
            print("\n⚠ 檢查未通過，請修正以上 FAIL 項目後再重新產出文件。")
        else:
            print("\n✓ 所有必要檢查通過！")


# .docx 檢查（使用 python-docx）
def _import_docx():
    try:
        from docx import Document
        return Document
    except ImportError:
        sys.exit("錯誤：檢查 .docx 需要 python-docx，請執行：pip install python-docx")


def check_styles(doc, result: InspectionResult) -> None:
    available_styles = {s.name for s in doc.styles if s.name is not None}
    print("\n--- 檢查樣式 ---")
    for style_name in REQUIRED_STYLES:
        if style_name in available_styles:
            result.add_pass(f"樣式存在：{style_name}")
        else:
            result.add_fail(f"缺少樣式：{style_name}")
    for style_name in LIST_STYLE_PREFIXES:
        if style_name in available_styles:
            result.add_pass(f"清單樣式存在：{style_name}")
        else:
            result.add_warn(f"缺少清單樣式：{style_name}")


def check_error_ref_not_found_docx(doc, result: InspectionResult) -> None:
    print("\n--- 檢查錯誤參照 ---")
    error_count = 0
    error_paragraphs: List[Tuple[int, str]] = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if "Error!" in text and ("Reference source not found" in text or "reference source" in text.lower()):
            error_count += 1
            snippet = text[:100] if len(text) > 100 else text
            error_paragraphs.append((i + 1, snippet))
    if error_count == 0:
        result.add_pass("沒有發現「Error: Reference source not found」錯誤")
    else:
        result.add_fail(f"發現 {error_count} 處「Error: Reference source not found」錯誤")
        for line_no, snippet in error_paragraphs[:10]:
            result.add_fail(f"  段落 {line_no}：{snippet}")


def check_normal_paragraph_ratio(doc, result: InspectionResult) -> None:
    print("\n--- 檢查 Normal 段落比例 ---")
    total_paras = 0
    normal_paras = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        total_paras += 1
        style_name = para.style.name if para.style else "None"
        if style_name == "Normal" or style_name == "normal":
            normal_paras += 1
    if total_paras == 0:
        result.add_warn("文件中沒有段落內容")
        return
    ratio = normal_paras / total_paras * 100
    if ratio > 50:
        result.add_fail(
            f"Normal 段落比例過高：{ratio:.1f}%（{normal_paras}/{total_paras}）。"
            f"請確認內容是否使用了正確的樣式（_內文-縮排、_內文-無縮排等）"
        )
    elif ratio > 30:
        result.add_warn(
            f"Normal 段落比例偏高：{ratio:.1f}%（{normal_paras}/{total_paras}）"
        )
    else:
        result.add_pass(f"Normal 段落比例正常：{ratio:.1f}%（{normal_paras}/{total_paras}）")


def check_cover_content(doc, result: InspectionResult) -> None:
    print("\n--- 檢查封面內容 ---")
    cover_text = ""
    for para in doc.paragraphs[:30]:
        cover_text += para.text + "\n"
    for keyword in COVER_REQUIRED:
        if keyword in cover_text:
            result.add_pass("封面包含 Template 版本標記")
        else:
            result.add_warn("封面可能缺少 Template 版本標記")
    if re.search(r'V\d+\.\d+\.\d+', cover_text):
        result.add_pass("封面包含版本號（如 V1.0.0）")
    else:
        result.add_warn("封面可能缺少版本號")
    if re.search(r'\d{4}[/-]\d{2}[/-]\d{2}', cover_text):
        result.add_pass("封面包含日期格式")
    else:
        result.add_warn("封面可能缺少日期")
    non_empty_paras = [p.text for p in doc.paragraphs[:15] if p.text.strip()]
    if len(non_empty_paras) >= 3:
        result.add_pass(f"封面有足夠內容（{len(non_empty_paras)} 個非空段落）")
    else:
        result.add_warn("封面內容可能不足")


def check_revision_table_docx(doc, result: InspectionResult) -> None:
    print("\n--- 檢查修訂紀錄表 ---")
    found_revision = False
    for para in doc.paragraphs:
        if "修訂紀錄" in para.text:
            found_revision = True
            break
    if found_revision:
        result.add_pass("找到修訂紀錄表")
    else:
        result.add_fail("找不到修訂紀錄表")
    if len(doc.tables) > 0:
        result.add_pass(f"文件中包含 {len(doc.tables)} 個表格")
    else:
        result.add_warn("文件中沒有任何表格")


def check_heading1(doc, result: InspectionResult) -> None:
    print("\n--- 檢查標題結構 ---")
    heading1_count = 0
    headings_found = set()
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            headings_found.add(style_name)
            if style_name == "Heading 1":
                heading1_count += 1
    if heading1_count > 0:
        result.add_pass(f"找到 {heading1_count} 個 Heading 1")
    else:
        result.add_fail("文件中沒有任何 Heading 1")
    if headings_found:
        result.add_pass(f"使用的標題層級：{', '.join(sorted(headings_found))}")
    else:
        result.add_warn("文件中沒有任何標題（Heading）")


def check_table_captions(doc, result: InspectionResult) -> None:
    print("\n--- 檢查表格標號 ---")
    caption_count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "表_標號":
            caption_count += 1
    if caption_count > 0:
        result.add_pass(f"找到 {caption_count} 個表標號")
    else:
        result.add_warn("未找到任何表標號")


def check_header_footer(doc, result: InspectionResult) -> None:
    print("\n--- 檢查頁首頁尾 ---")
    has_header = False
    has_footer = False
    for section in doc.sections:
        if section.header and not section.header.is_linked_to_previous:
            has_header = True
        if section.footer and not section.footer.is_linked_to_previous:
            has_footer = True
    if has_header:
        result.add_pass("頁首存在")
    else:
        result.add_warn("頁首可能不存在或未正確設定")
    if has_footer:
        result.add_pass("頁尾存在")
    else:
        result.add_warn("頁尾可能不存在或未正確設定")


def inspect_docx(filepath: str) -> InspectionResult:
    path = Path(filepath)
    if not path.exists():
        result = InspectionResult()
        result.add_fail(f"檔案不存在：{filepath}")
        return result
    print(f"檢查檔案：{path.resolve()}")
    print(f"檔案大小：{path.stat().st_size / 1024:.1f} KB")
    print(f"檔案格式：.docx（Office Open XML）")
    Document = _import_docx()
    doc = Document(str(path))
    result = InspectionResult()
    check_styles(doc, result)
    check_error_ref_not_found_docx(doc, result)
    check_normal_paragraph_ratio(doc, result)
    check_cover_content(doc, result)
    check_revision_table_docx(doc, result)
    check_heading1(doc, result)
    check_table_captions(doc, result)
    check_header_footer(doc, result)
    return result


# .doc 檢查（使用 olefile）
def _import_olefile():
    try:
        import olefile
        return olefile
    except ImportError:
        # 嘗試從專案 venv 載入
        _venv_path = Path(__file__).resolve().parent / '.venv' / 'Lib' / 'site-packages'
        if _venv_path.exists():
            sys.path.insert(0, str(_venv_path))
            try:
                import olefile
                return olefile
            except ImportError:
                pass
        sys.exit("錯誤：檢查 .doc 需要 olefile，請執行：pip install olefile")


def _extract_text_from_doc(filepath: str) -> str:
    olefile = _import_olefile()
    ole = olefile.OleFileIO(filepath)
    word_stream = ole.openstream('WordDocument').read()
    ole.close()
    result = []
    text_buffer = []
    i = 0
    while i < len(word_stream) - 1:
        code_unit = word_stream[i] | (word_stream[i + 1] << 8)
        if ((0x20 <= code_unit <= 0x7E) or
                (0x4E00 <= code_unit <= 0x9FFF) or
                (0x3400 <= code_unit <= 0x4DBF) or
                (0x3000 <= code_unit <= 0x303F) or
                (0xFF00 <= code_unit <= 0xFFEF) or
                (0x2000 <= code_unit <= 0x206F) or
                code_unit in (0x0D, 0x0A, 0x09)):
            text_buffer.append(chr(code_unit))
        else:
            if len(text_buffer) >= 3:
                result.append(''.join(text_buffer))
            text_buffer = []
        i += 2
    if len(text_buffer) >= 3:
        result.append(''.join(text_buffer))
    return '\n'.join(result)


def _check_error_ref_in_text(text: str, result: InspectionResult) -> None:
    print("\n--- 檢查錯誤參照 ---")
    pattern = re.compile(r'Error!.*?[Rr]eference\s+source\s+not\s+found', re.IGNORECASE)
    matches = pattern.findall(text)
    if not matches:
        result.add_pass("沒有發現「Error: Reference source not found」錯誤")
    else:
        result.add_fail(f"發現 {len(matches)} 處「Error: Reference source not found」錯誤")
        for m in matches[:5]:
            result.add_fail(f"  → {m[:80]}")


def _check_revision_in_text(text: str, result: InspectionResult) -> None:
    print("\n--- 檢查修訂紀錄表 ---")
    if "修訂紀錄" in text:
        result.add_pass("找到修訂紀錄")
    else:
        result.add_fail("找不到修訂紀錄")


def _check_content_structure(text: str, result: InspectionResult) -> None:
    print("\n--- 檢查內容結構 ---")
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if len(lines) > 0:
        result.add_pass(f"文件包含 {len(lines)} 行有效文字內容")
    else:
        result.add_fail("文件沒有可讀取的文字內容")
        return
    if re.search(r'V\d+\.\d+\.\d+', text):
        result.add_pass("文件中包含版本號（如 V1.0.0）")
    else:
        result.add_warn("文件中可能缺少版本號")
    if re.search(r'\d{4}[/-]\d{2}[/-]\d{2}', text):
        result.add_pass("文件中包含日期格式")
    else:
        result.add_warn("文件中可能缺少日期")
    if "Template" in text:
        result.add_pass("文件中包含 Template 標記")
    else:
        result.add_warn("文件中可能缺少 Template 版本標記")


def _check_doc_basic_format(text: str, result: InspectionResult) -> None:
    print("\n--- 檢查基本格式 ---")
    paragraphs = [p for p in text.split('\n') if p.strip()]
    if len(paragraphs) >= 5:
        result.add_pass(f"文件結構合理（{len(paragraphs)} 個段落）")
    else:
        result.add_warn("文件段落數量偏少，可能結構不完整")
    short_lines = [p for p in paragraphs if 3 <= len(p) <= 60]
    if short_lines:
        result.add_pass(f"發現 {len(short_lines)} 行疑似標題的短文字")


def inspect_doc(filepath: str) -> InspectionResult:
    path = Path(filepath)
    if not path.exists():
        result = InspectionResult()
        result.add_fail(f"檔案不存在：{filepath}")
        return result
    print(f"檢查檔案：{path.resolve()}")
    print(f"檔案大小：{path.stat().st_size / 1024:.1f} KB")
    print(f"檔案格式：.doc（OLE Compound File）")
    print("注意：.doc 格式僅支援基本內容檢查，樣式檢查需先轉為 .docx")
    result = InspectionResult()
    try:
        text = _extract_text_from_doc(filepath)
    except Exception as e:
        result.add_fail(f"無法讀取 .doc 檔案內容：{e}")
        return result
    if not text.strip():
        result.add_fail("無法從 .doc 檔案中擷取到任何文字內容")
        return result
    result.add_pass(f"成功擷取文字內容（{len(text)} 字元）")
    _check_error_ref_in_text(text, result)
    _check_revision_in_text(text, result)
    _check_content_structure(text, result)
    _check_doc_basic_format(text, result)
    return result


# 主入口
def main() -> None:
    parser = argparse.ArgumentParser(
        description="檢查公司格式 Word 文件是否符合樣式規範（支援 .docx 與 .doc）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
範例：
    python inspect_docx.py output.docx
    python inspect_docx.py CM03A_V1.doc
    python inspect_docx.py demo.docx

支援格式：
    .docx  → 完整檢查：樣式、錯誤參照、段落、封面、修訂紀錄、標題、表格、頁首頁尾
    .doc   → 基本檢查：錯誤參照、修訂紀錄、內容結構、基本格式
             （樣式檢查需先將 .doc 轉為 .docx）
        """,
    )
    parser.add_argument(
        "filepath",
        help="要檢查的 Word 檔案路徑（.docx 或 .doc）",
    )
    args = parser.parse_args()

    file_type = _get_file_type(args.filepath)

    if file_type == 'docx':
        result = inspect_docx(args.filepath)
    elif file_type == 'doc':
        result = inspect_doc(args.filepath)
    else:
        print(f"錯誤：不支援的檔案格式「{Path(args.filepath).suffix}」，僅支援 .docx 與 .doc")
        sys.exit(1)

    result.print_summary()

    if result.has_failures():
        sys.exit(1)


if __name__ == "__main__":
    main()